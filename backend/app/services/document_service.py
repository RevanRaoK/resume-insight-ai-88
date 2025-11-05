"""
Document ingestion service for processing various document formats
"""
import os
import tempfile
from typing import Optional, Dict, Any, AsyncGenerator
from abc import ABC, abstractmethod
from fastapi import UploadFile

import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import chardet

from app.utils.logger import get_logger
from app.utils.async_utils import async_timer, async_pipeline
from app.core.exceptions import DocumentProcessingError, UnsupportedFormatError
from app.models.entities import ProcessedDocument
from app.utils.file_utils import validate_upload_file, create_secure_temp_file
from app.core.security import TemporaryFileManager

logger = get_logger(__name__)


class BaseDocumentProcessor(ABC):
    """Abstract base class for document processors"""
    
    @abstractmethod
    @async_timer
    async def process(self, file_path: str, filename: str) -> ProcessedDocument:
        """Process document and extract text"""
        pass
    
    @abstractmethod
    def supports_format(self, mime_type: str) -> bool:
        """Check if processor supports the given MIME type"""
        pass


class PDFProcessor(BaseDocumentProcessor):
    """PDF document processor using pdfplumber for digital PDFs"""
    
    def supports_format(self, mime_type: str) -> bool:
        return mime_type == "application/pdf"
    
    async def process(self, file_path: str, filename: str) -> ProcessedDocument:
        """
        Extract text from PDF using pdfplumber
        
        Args:
            file_path: Path to the PDF file
            filename: Original filename
            
        Returns:
            ProcessedDocument with extracted text
            
        Raises:
            DocumentProcessingError: If PDF processing fails
        """
        try:
            logger.info("pdf_processing_started", filename=filename)
            
            extracted_text = ""
            page_count = 0
            
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            extracted_text += page_text + "\n"
                        
                        logger.debug(
                            "pdf_page_processed",
                            filename=filename,
                            page_number=page_num,
                            text_length=len(page_text) if page_text else 0
                        )
                    except Exception as e:
                        logger.warning(
                            "pdf_page_processing_failed",
                            filename=filename,
                            page_number=page_num,
                            error=str(e)
                        )
                        continue
            
            # Calculate confidence based on text length and page coverage
            confidence_score = min(1.0, len(extracted_text.strip()) / (page_count * 200))
            
            file_size = os.path.getsize(file_path)
            
            logger.info(
                "pdf_processing_completed",
                filename=filename,
                pages_processed=page_count,
                text_length=len(extracted_text),
                confidence_score=confidence_score
            )
            
            return ProcessedDocument(
                text=extracted_text.strip(),
                file_name=filename,
                file_size=file_size,
                processing_method="pdfplumber",
                confidence_score=confidence_score
            )
            
        except Exception as e:
            logger.error(
                "pdf_processing_error",
                filename=filename,
                error=str(e),
                error_type=type(e).__name__
            )
            raise DocumentProcessingError(
                message=f"Failed to process PDF: {str(e)}",
                file_name=filename,
                processing_stage="pdf_extraction"
            )


class OCRProcessor(BaseDocumentProcessor):
    """OCR processor for scanned PDFs using pdf2image and pytesseract"""
    
    def supports_format(self, mime_type: str) -> bool:
        return mime_type == "application/pdf"
    
    async def process(self, file_path: str, filename: str) -> ProcessedDocument:
        """
        Extract text from scanned PDF using OCR
        
        Args:
            file_path: Path to the PDF file
            filename: Original filename
            
        Returns:
            ProcessedDocument with OCR-extracted text
            
        Raises:
            DocumentProcessingError: If OCR processing fails
        """
        try:
            logger.info("ocr_processing_started", filename=filename)
            
            # Convert PDF pages to images
            images = convert_from_path(file_path, dpi=300)
            
            extracted_text = ""
            
            for page_num, image in enumerate(images, 1):
                try:
                    # Perform OCR on the image
                    page_text = pytesseract.image_to_string(
                        image, 
                        config='--psm 6 -l eng'  # Page segmentation mode 6, English language
                    )
                    
                    if page_text.strip():
                        extracted_text += page_text + "\n"
                    
                    logger.debug(
                        "ocr_page_processed",
                        filename=filename,
                        page_number=page_num,
                        text_length=len(page_text)
                    )
                    
                except Exception as e:
                    logger.warning(
                        "ocr_page_processing_failed",
                        filename=filename,
                        page_number=page_num,
                        error=str(e)
                    )
                    continue
            
            # OCR confidence is generally lower than digital text extraction
            confidence_score = min(0.8, len(extracted_text.strip()) / (len(images) * 150))
            
            file_size = os.path.getsize(file_path)
            
            logger.info(
                "ocr_processing_completed",
                filename=filename,
                pages_processed=len(images),
                text_length=len(extracted_text),
                confidence_score=confidence_score
            )
            
            return ProcessedDocument(
                text=extracted_text.strip(),
                file_name=filename,
                file_size=file_size,
                processing_method="ocr",
                confidence_score=confidence_score
            )
            
        except Exception as e:
            logger.error(
                "ocr_processing_error",
                filename=filename,
                error=str(e),
                error_type=type(e).__name__
            )
            raise DocumentProcessingError(
                message=f"Failed to process PDF with OCR: {str(e)}",
                file_name=filename,
                processing_stage="ocr_extraction"
            )


class DOCXProcessor(BaseDocumentProcessor):
    """DOCX document processor using python-docx"""
    
    def supports_format(self, mime_type: str) -> bool:
        return mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    async def process(self, file_path: str, filename: str) -> ProcessedDocument:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to the DOCX file
            filename: Original filename
            
        Returns:
            ProcessedDocument with extracted text
            
        Raises:
            DocumentProcessingError: If DOCX processing fails
        """
        try:
            logger.info("docx_processing_started", filename=filename)
            
            doc = Document(file_path)
            
            extracted_text = ""
            paragraph_count = 0
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_text += paragraph.text + "\n"
                    paragraph_count += 1
            
            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        extracted_text += " | ".join(row_text) + "\n"
            
            # High confidence for DOCX as it's structured text
            confidence_score = 0.95 if extracted_text.strip() else 0.0
            
            file_size = os.path.getsize(file_path)
            
            logger.info(
                "docx_processing_completed",
                filename=filename,
                paragraphs_processed=paragraph_count,
                tables_processed=table_count,
                text_length=len(extracted_text),
                confidence_score=confidence_score
            )
            
            return ProcessedDocument(
                text=extracted_text.strip(),
                file_name=filename,
                file_size=file_size,
                processing_method="docx",
                confidence_score=confidence_score
            )
            
        except Exception as e:
            logger.error(
                "docx_processing_error",
                filename=filename,
                error=str(e),
                error_type=type(e).__name__
            )
            raise DocumentProcessingError(
                message=f"Failed to process DOCX: {str(e)}",
                file_name=filename,
                processing_stage="docx_extraction"
            )


class TextProcessor(BaseDocumentProcessor):
    """Plain text processor with encoding detection"""
    
    def supports_format(self, mime_type: str) -> bool:
        return mime_type == "text/plain"
    
    async def process(self, file_path: str, filename: str) -> ProcessedDocument:
        """
        Process plain text file with encoding detection
        
        Args:
            file_path: Path to the text file
            filename: Original filename
            
        Returns:
            ProcessedDocument with extracted text
            
        Raises:
            DocumentProcessingError: If text processing fails
        """
        try:
            logger.info("text_processing_started", filename=filename)
            
            # Read file in binary mode for encoding detection
            with open(file_path, 'rb') as file:
                raw_data = file.read()
            
            # Detect encoding
            encoding_result = chardet.detect(raw_data)
            detected_encoding = encoding_result.get('encoding', 'utf-8')
            encoding_confidence = encoding_result.get('confidence', 0.0)
            
            logger.debug(
                "encoding_detected",
                filename=filename,
                detected_encoding=detected_encoding,
                encoding_confidence=encoding_confidence
            )
            
            # Decode text with detected encoding
            try:
                extracted_text = raw_data.decode(detected_encoding)
            except (UnicodeDecodeError, TypeError, LookupError):
                # Fallback to utf-8 with error handling
                logger.warning(
                    "encoding_fallback",
                    filename=filename,
                    failed_encoding=detected_encoding
                )
                extracted_text = raw_data.decode('utf-8', errors='replace')
                encoding_confidence = 0.5
            
            # Confidence based on encoding detection and text quality
            confidence_score = min(0.95, encoding_confidence + 0.1)
            
            file_size = os.path.getsize(file_path)
            
            logger.info(
                "text_processing_completed",
                filename=filename,
                encoding_used=detected_encoding,
                text_length=len(extracted_text),
                confidence_score=confidence_score
            )
            
            return ProcessedDocument(
                text=extracted_text.strip(),
                file_name=filename,
                file_size=file_size,
                processing_method="text",
                confidence_score=confidence_score
            )
            
        except Exception as e:
            logger.error(
                "text_processing_error",
                filename=filename,
                error=str(e),
                error_type=type(e).__name__
            )
            raise DocumentProcessingError(
                message=f"Failed to process text file: {str(e)}",
                file_name=filename,
                processing_stage="text_extraction"
            )


class DocumentService:
    """Main document service that orchestrates different processors"""
    
    def __init__(self):
        self.processors = {
            "application/pdf": [PDFProcessor(), OCRProcessor()],  # PDF with OCR fallback
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": [DOCXProcessor()],
            "text/plain": [TextProcessor()]
        }
    
    async def process_document(self, upload_file: UploadFile) -> ProcessedDocument:
        """
        Process uploaded document using appropriate processor
        
        Args:
            upload_file: FastAPI UploadFile object
            
        Returns:
            ProcessedDocument with extracted text and metadata
            
        Raises:
            UnsupportedFormatError: If file format is not supported
            DocumentProcessingError: If processing fails
        """
        # Validate file first
        validation_result = await validate_upload_file(upload_file)
        
        mime_type = validation_result["detected_mime_type"]
        safe_filename = validation_result["safe_filename"]
        content = validation_result["content"]
        
        # Check if format is supported
        if mime_type not in self.processors:
            supported_types = list(self.processors.keys())
            raise UnsupportedFormatError(
                file_type=mime_type,
                supported_types=supported_types
            )
        
        # Create temporary file for processing with automatic cleanup
        with TemporaryFileManager() as temp_manager:
            # Extract file extension for proper handling
            _, ext = os.path.splitext(safe_filename)
            temp_file_path = temp_manager.create_temp_file(content, ext)
            processors = self.processors[mime_type]
            last_error = None
            
            # Try processors in order (for PDF: pdfplumber first, then OCR fallback)
            for processor in processors:
                try:
                    logger.info(
                        "document_processing_attempt",
                        filename=safe_filename,
                        processor=processor.__class__.__name__,
                        mime_type=mime_type
                    )
                    
                    result = await processor.process(temp_file_path, safe_filename)
                    
                    # For PDF, check if we need OCR fallback
                    if (mime_type == "application/pdf" and 
                        isinstance(processor, PDFProcessor) and 
                        len(result.text.strip()) < 200):
                        
                        logger.info(
                            "pdf_ocr_fallback_triggered",
                            filename=safe_filename,
                            extracted_text_length=len(result.text.strip())
                        )
                        continue  # Try OCR processor
                    
                    logger.info(
                        "document_processing_successful",
                        filename=safe_filename,
                        processor=processor.__class__.__name__,
                        text_length=len(result.text),
                        confidence_score=result.confidence_score
                    )
                    
                    return result
                    
                except Exception as e:
                    last_error = e
                    logger.warning(
                        "processor_failed",
                        filename=safe_filename,
                        processor=processor.__class__.__name__,
                        error=str(e)
                    )
                    continue
            
            # If all processors failed, raise the last error
            if last_error:
                raise last_error
            else:
                raise DocumentProcessingError(
                    message="No suitable processor found",
                    file_name=safe_filename,
                    processing_stage="processor_selection"
                )
            # Temporary file cleanup is handled automatically by TemporaryFileManager context manager
    
    async def process_documents_stream(
        self, 
        upload_files: list[UploadFile]
    ) -> AsyncGenerator[ProcessedDocument, None]:
        """
        Process multiple documents using async streaming pipeline
        
        Args:
            upload_files: List of FastAPI UploadFile objects
            
        Yields:
            ProcessedDocument objects as they are processed
        """
        logger.info(
            "streaming_document_processing_started",
            total_files=len(upload_files)
        )
        
        async def process_single_file(upload_file: UploadFile) -> ProcessedDocument:
            """Process a single file for the streaming pipeline"""
            return await self.process_document(upload_file)
        
        # Use async pipeline for streaming processing
        async for result in async_pipeline.process_in_chunks(
            upload_files,
            process_single_file,
            chunk_size=3  # Process 3 files concurrently
        ):
            yield result
        
        logger.info(
            "streaming_document_processing_completed",
            total_files=len(upload_files)
        )