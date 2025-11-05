"""
Unit tests for individual document processors
"""
import pytest
import os
from unittest.mock import patch, Mock, MagicMock

from app.services.document_service import (
    PDFProcessor, 
    OCRProcessor, 
    DOCXProcessor, 
    TextProcessor
)
from app.core.exceptions import DocumentProcessingError
from app.models.entities import ProcessedDocument


class TestPDFProcessor:
    """Test cases for PDFProcessor"""
    
    @pytest.fixture
    def pdf_processor(self):
        return PDFProcessor()
    
    def test_supports_format(self, pdf_processor):
        """Test MIME type support detection"""
        assert pdf_processor.supports_format("application/pdf") is True
        assert pdf_processor.supports_format("text/plain") is False
        assert pdf_processor.supports_format("application/vnd.openxmlformats-officedocument.wordprocessingml.document") is False
    
    @pytest.mark.asyncio
    async def test_process_valid_pdf(self, pdf_processor, create_test_txt, sample_text_content):
        """Test processing a valid PDF file"""
        # Use text file as mock PDF since creating real PDF is complex
        txt_path = create_test_txt
        
        # Mock pdfplumber to return expected text
        with patch('app.services.document_service.pdfplumber') as mock_pdfplumber:
            mock_pdf = Mock()
            mock_page = Mock()
            mock_page.extract_text.return_value = sample_text_content
            mock_pdf.pages = [mock_page]
            mock_pdfplumber.open.return_value.__enter__.return_value = mock_pdf
            
            result = await pdf_processor.process(txt_path, "test.pdf")
            
            assert isinstance(result, ProcessedDocument)
            assert result.text == sample_text_content.strip().strip()
            assert result.file_name == "test.pdf"
            assert result.processing_method == "pdfplumber"
            assert result.confidence_score > 0
            assert result.file_size > 0
    
    @pytest.mark.asyncio
    async def test_process_pdf_with_empty_pages(self, pdf_processor, create_test_txt):
        """Test processing PDF with some empty pages"""
        txt_path = create_test_txt
        
        with patch('app.services.document_service.pdfplumber') as mock_pdfplumber:
            mock_pdf = Mock()
            mock_page1 = Mock()
            mock_page1.extract_text.return_value = "Page 1 content"
            mock_page2 = Mock()
            mock_page2.extract_text.return_value = None  # Empty page
            mock_page3 = Mock()
            mock_page3.extract_text.return_value = "Page 3 content"
            
            mock_pdf.pages = [mock_page1, mock_page2, mock_page3]
            mock_pdfplumber.open.return_value.__enter__.return_value = mock_pdf
            
            result = await pdf_processor.process(txt_path, "test.pdf")
            
            assert "Page 1 content" in result.text
            assert "Page 3 content" in result.text
            assert result.processing_method == "pdfplumber"
    
    @pytest.mark.asyncio
    async def test_process_pdf_extraction_error(self, pdf_processor, create_test_txt):
        """Test handling of PDF extraction errors"""
        txt_path = create_test_txt
        
        with patch('app.services.document_service.pdfplumber') as mock_pdfplumber:
            mock_pdfplumber.open.side_effect = Exception("PDF parsing failed")
            
            with pytest.raises(DocumentProcessingError) as exc_info:
                await pdf_processor.process(txt_path, "test.pdf")
            
            assert "Failed to process PDF" in str(exc_info.value)
            assert exc_info.value.details["file_name"] == "test.pdf"
            assert exc_info.value.details["processing_stage"] == "pdf_extraction"


class TestOCRProcessor:
    """Test cases for OCRProcessor"""
    
    @pytest.fixture
    def ocr_processor(self):
        return OCRProcessor()
    
    def test_supports_format(self, ocr_processor):
        """Test MIME type support detection"""
        assert ocr_processor.supports_format("application/pdf") is True
        assert ocr_processor.supports_format("text/plain") is False
    
    @pytest.mark.asyncio
    async def test_process_scanned_pdf(self, ocr_processor, create_test_txt, sample_text_content):
        """Test OCR processing of scanned PDF"""
        txt_path = create_test_txt
        
        # Mock pdf2image and pytesseract
        with patch('app.services.document_service.convert_from_path') as mock_convert, \
             patch('app.services.document_service.pytesseract') as mock_tesseract:
            
            # Mock image conversion
            mock_image = Mock()
            mock_convert.return_value = [mock_image, mock_image]  # 2 pages
            
            # Mock OCR text extraction
            mock_tesseract.image_to_string.return_value = sample_text_content[:100]
            
            result = await ocr_processor.process(txt_path, "scanned.pdf")
            
            assert isinstance(result, ProcessedDocument)
            assert len(result.text) > 0
            assert result.file_name == "scanned.pdf"
            assert result.processing_method == "ocr"
            assert result.confidence_score <= 0.8  # OCR has lower confidence
            
            # Verify OCR was called with correct parameters
            mock_tesseract.image_to_string.assert_called_with(
                mock_image, 
                config='--psm 6 -l eng'
            )
    
    @pytest.mark.asyncio
    async def test_process_ocr_with_failed_pages(self, ocr_processor, create_test_txt):
        """Test OCR processing with some failed pages"""
        txt_path = create_test_txt
        
        with patch('app.services.document_service.convert_from_path') as mock_convert, \
             patch('app.services.document_service.pytesseract') as mock_tesseract:
            
            mock_image1 = Mock()
            mock_image2 = Mock()
            mock_convert.return_value = [mock_image1, mock_image2]
            
            # First page succeeds, second fails
            mock_tesseract.image_to_string.side_effect = ["Page 1 text", Exception("OCR failed")]
            
            result = await ocr_processor.process(txt_path, "scanned.pdf")
            
            assert "Page 1 text" in result.text
            assert result.processing_method == "ocr"
    
    @pytest.mark.asyncio
    async def test_process_ocr_conversion_error(self, ocr_processor, create_test_txt):
        """Test handling of PDF to image conversion errors"""
        txt_path = create_test_txt
        
        with patch('app.services.document_service.convert_from_path') as mock_convert:
            mock_convert.side_effect = Exception("PDF conversion failed")
            
            with pytest.raises(DocumentProcessingError) as exc_info:
                await ocr_processor.process(txt_path, "scanned.pdf")
            
            assert "Failed to process PDF with OCR" in str(exc_info.value)
            assert exc_info.value.details["processing_stage"] == "ocr_extraction"


class TestDOCXProcessor:
    """Test cases for DOCXProcessor"""
    
    @pytest.fixture
    def docx_processor(self):
        return DOCXProcessor()
    
    def test_supports_format(self, docx_processor):
        """Test MIME type support detection"""
        assert docx_processor.supports_format("application/vnd.openxmlformats-officedocument.wordprocessingml.document") is True
        assert docx_processor.supports_format("application/pdf") is False
        assert docx_processor.supports_format("text/plain") is False
    
    @pytest.mark.asyncio
    async def test_process_valid_docx(self, docx_processor, create_test_docx, sample_text_content):
        """Test processing a valid DOCX file"""
        docx_path = create_test_docx
        
        result = await docx_processor.process(docx_path, "test.docx")
        
        assert isinstance(result, ProcessedDocument)
        assert len(result.text) > 0
        assert result.file_name == "test.docx"
        assert result.processing_method == "docx"
        assert result.confidence_score == 0.95  # High confidence for structured text
        assert result.file_size > 0
    
    @pytest.mark.asyncio
    async def test_process_docx_with_tables(self, docx_processor, temp_dir):
        """Test processing DOCX with tables"""
        from docx import Document
        
        # Create DOCX with table
        doc = Document()
        doc.add_paragraph("Header text")
        
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "John Doe"
        table.cell(1, 0).text = "Position"
        table.cell(1, 1).text = "Software Engineer"
        
        docx_path = os.path.join(temp_dir, "test_with_table.docx")
        doc.save(docx_path)
        
        result = await docx_processor.process(docx_path, "test_with_table.docx")
        
        assert "Header text" in result.text
        assert "Name | John Doe" in result.text
        assert "Position | Software Engineer" in result.text
        assert result.processing_method == "docx"
    
    @pytest.mark.asyncio
    async def test_process_empty_docx(self, docx_processor, temp_dir):
        """Test processing empty DOCX file"""
        from docx import Document
        
        doc = Document()
        docx_path = os.path.join(temp_dir, "empty.docx")
        doc.save(docx_path)
        
        result = await docx_processor.process(docx_path, "empty.docx")
        
        assert result.text == ""
        assert result.confidence_score == 0.0
        assert result.processing_method == "docx"
    
    @pytest.mark.asyncio
    async def test_process_corrupted_docx(self, docx_processor, create_corrupted_pdf):
        """Test handling of corrupted DOCX files"""
        corrupted_path = create_corrupted_pdf  # Use corrupted file
        
        with pytest.raises(DocumentProcessingError) as exc_info:
            await docx_processor.process(corrupted_path, "corrupted.docx")
        
        assert "Failed to process DOCX" in str(exc_info.value)
        assert exc_info.value.details["processing_stage"] == "docx_extraction"


class TestTextProcessor:
    """Test cases for TextProcessor"""
    
    @pytest.fixture
    def text_processor(self):
        return TextProcessor()
    
    def test_supports_format(self, text_processor):
        """Test MIME type support detection"""
        assert text_processor.supports_format("text/plain") is True
        assert text_processor.supports_format("application/pdf") is False
        assert text_processor.supports_format("application/vnd.openxmlformats-officedocument.wordprocessingml.document") is False
    
    @pytest.mark.asyncio
    async def test_process_utf8_text(self, text_processor, create_test_txt, sample_text_content):
        """Test processing UTF-8 text file"""
        txt_path = create_test_txt
        
        result = await text_processor.process(txt_path, "test.txt")
        
        assert isinstance(result, ProcessedDocument)
        # Normalize line endings for cross-platform compatibility
        expected_text = sample_text_content.strip().replace('\r\n', '\n')
        actual_text = result.text.replace('\r\n', '\n')
        assert actual_text == expected_text
        assert result.file_name == "test.txt"
        assert result.processing_method == "text"
        assert result.confidence_score >= 0.9  # High confidence for text files
        assert result.file_size > 0
    
    @pytest.mark.asyncio
    async def test_process_different_encoding(self, text_processor, temp_dir):
        """Test processing text file with different encoding"""
        # Create file with Latin-1 encoding
        content = "Résumé with special characters: café, naïve"
        txt_path = os.path.join(temp_dir, "latin1.txt")
        
        with open(txt_path, "w", encoding="latin-1") as f:
            f.write(content)
        
        result = await text_processor.process(txt_path, "latin1.txt")
        
        assert len(result.text) > 0
        assert result.processing_method == "text"
        # Should handle encoding detection
    
    @pytest.mark.asyncio
    async def test_process_encoding_fallback(self, text_processor, temp_dir):
        """Test encoding fallback mechanism"""
        # Create file with mixed/invalid encoding
        txt_path = os.path.join(temp_dir, "mixed_encoding.txt")
        
        with open(txt_path, "wb") as f:
            f.write(b"Valid text\xff\xfe\x00Invalid bytes")
        
        # Mock chardet to return unreliable encoding
        with patch('app.services.document_service.chardet') as mock_chardet:
            mock_chardet.detect.return_value = {
                'encoding': 'invalid-encoding',
                'confidence': 0.3
            }
            
            result = await text_processor.process(txt_path, "mixed_encoding.txt")
            
            assert len(result.text) > 0  # Should fallback to UTF-8 with error handling
            assert result.processing_method == "text"
            assert result.confidence_score <= 0.6  # Lower confidence due to fallback
    
    @pytest.mark.asyncio
    async def test_process_binary_file_as_text(self, text_processor, create_unsupported_file):
        """Test processing binary file as text (should handle gracefully)"""
        binary_path = create_unsupported_file
        
        result = await text_processor.process(binary_path, "binary.txt")
        
        # Should not crash, but may have replacement characters
        assert isinstance(result, ProcessedDocument)
        assert result.processing_method == "text"
        # Confidence should be low due to binary content
    
    @pytest.mark.asyncio
    async def test_process_file_read_error(self, text_processor):
        """Test handling of file read errors"""
        non_existent_path = "/non/existent/file.txt"
        
        with pytest.raises(DocumentProcessingError) as exc_info:
            await text_processor.process(non_existent_path, "missing.txt")
        
        assert "Failed to process text file" in str(exc_info.value)
        assert exc_info.value.details["processing_stage"] == "text_extraction"